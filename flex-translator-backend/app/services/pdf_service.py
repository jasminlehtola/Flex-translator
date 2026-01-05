"""
pdf_service.py

Utility functions for working with PDF and DOCX files, including:
- Converting PDFs to DOCX
- Extracting text from DOCX
- Applying translations to DOCX content
- Inferring file types from binary data
"""

import base64
from io import BytesIO
from pdf2docx import Converter
from docx import Document
import zipfile 
import os
from lxml import etree
import tempfile
import zipfile

def extract_text(stream: BytesIO):
    """
    Extracts all text content from the DOCX (XML) stream.

    Args:
        stream (BytesIO): A file-like object of the DOCX file.

    Returns:
        str: All text from the document.
    """
    text = []
    with zipfile.ZipFile(stream) as docx:
        for file_name in docx.namelist():
            if file_name == "word/document.xml":#file_name.endswith(".xml"):
                with docx.open(file_name) as file:
                    tree = etree.parse(file)
                    try:
                        for elem in tree.iter():
                            if elem.text and elem.text.strip():
                                text.append(elem.text.strip())
                            #if elem.tail and elem.tail.strip():
                            #    text.append(elem.tail.strip())
                    except Exception:
                        print("Problem with parsing the text.")
    return  "\n".join(text)


def translate_text(docx_stream: str, translations: dict):
    """
    Applies translations to DOCX by modifying `docx.Document` paragraphs and headers.

    Args:
        docx_stream_b64 (str): Base64-encoded DOCX document.
        translations (dict): Mapping of original -> translated terms.

    Returns:
        BytesIO: Modified DOCX stream.
    """
    docx_stream = BytesIO(base64.b64decode(docx_stream[3:]))
    output_stream = BytesIO()
    translation_array = [(old, new) for (old,new) in translations.items()]

    def sorting_key(x):
        return len(x[0])
    translation_array = sorted(translation_array, key=sorting_key, reverse=True)
    with zipfile.ZipFile(docx_stream) as zip_ref:
        in_memory_files = {}

        for item in zip_ref.infolist():
            data = zip_ref.read(item.filename)

            if item.filename == "word/document.xml" :#item.filename.endswith(".xml"):
                try:
                    xml_tree = etree.fromstring(data)
                    for node in xml_tree.iter():
                        if node.text:
                            for (old, new) in translation_array:
                                if old in node.text:
                                    node.text = node.text.replace(old, new)
                                    """
                        if node.tail:
                            for old, new in translations.items():
                                if old in node.tail:
                                    node.tail = node.tail.replace(old, new) 
                                    """
                    new_data = etree.tostring(xml_tree, xml_declaration=True, encoding="UTF-8")

                    in_memory_files[item.filename] = new_data
                except Exception:
                    in_memory_files[item.filename] = data
            else:
                in_memory_files[item.filename] = data

    with zipfile.ZipFile(output_stream, "w", zipfile.ZIP_DEFLATED) as out:
        for name, data in in_memory_files.items():
            out.writestr(name, data)

    output_stream.seek(0)
    return output_stream


def translate_text2(docx_stream: str, translations: dict):
    """
    Applies word replacements inside a base64-encoded DOCX document (by directly editing XML).

    Args:
        docx_stream_b64 (str): Base64-encoded DOCX document, prefixed with "b64".
        translations (dict): Mapping of original -> translated terms.

    Returns:
        BytesIO: Translated DOCX as binary stream.
    """
    docx_stream = BytesIO(base64.b64decode(docx_stream[3:]))
    output_stream = BytesIO()
    translation_array = [(old, new) for (old,new) in translations.items()]
    doc = Document(docx_stream)

    def sorting_key(x):
        return len(x[0])
    translation_array = sorted(translation_array, key=sorting_key, reverse=True)
    
    for para in doc.paragraphs:
        for run in para.runs:
            for (old, new) in translation_array:
                if old in run.text:
                    print(run.text)
                    run.text = run.text.replace(old, new)

    for section in doc.sections:
        for header in [section.header, section.footer]:
            for para in header.paragraphs:
                for run in para.runs:
                    for (old, new) in translation_array:
                        if old in run.text:
                            run.text = run.text.replace(old, new)


    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream


def pdf2docx_encode_b64(file: str):
    """
    Converts a base64-encoded PDF (with "b64" prefix) into a DOCX file and re-encodes it to base64.

    Args:
        b64_pdf (str): Base64-encoded PDF string.

    Returns:
        str: Base64-encoded DOCX string with "b64" prefix.
    """
    pdf_bytes = base64.b64decode(file[3:])
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_pdf.write(pdf_bytes)
        tmp_pdf_path = tmp_pdf.name

    with  tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        tmp_docx_path = tmp_docx.name

    cv = Converter(tmp_pdf_path)
    cv.convert(tmp_docx_path)
    cv.close()

    with open(tmp_docx_path, "rb") as f:
        docx_bytes = base64.b64encode(f.read())
    docx_bytes = "b64"+docx_bytes.decode("utf-8")

    os.remove(tmp_pdf_path)
    os.remove(tmp_docx_path)
    return docx_bytes


def extract_text_from_pdf(text: str):
    """
    Extracts readable text from a base64-encoded PDF or returns plain text as-is.

    Args:
        text_or_b64 (str): Either plain text or base64-encoded PDF with prefix.

    Returns:
        str: Extracted or original text with <word> tags added.
    """
    if text[0:3] != "b64":
        return text
    
    pdf_bytes = base64.b64decode(text[3:])
    pdf_stream = BytesIO(pdf_bytes)
    output_text = extract_text(pdf_stream)
    output_text = output_text.split("\n")
    new_output = []
    for j in range(len(output_text)):
        # Remove all special characters from the end and beginning
        if sum([0 if not letter.isalpha() else 1 for letter in output_text[j]])==0:
               continue
        new_output.append(output_text[j])
    output_text = "\n".join(["<word>"+i+"<word>" for i in new_output])

    return output_text



# ---- DeepL-specific helpers ----

def guess_extension(file_bytes: bytes) -> str:
    """
    Guesses file extension from binary data using magic bytes and structure.

    Args:
        file_bytes (bytes): Raw file bytes.

    Returns:
        str: Inferred file extension (.pdf, .docx, .pptx, etc.)
    """
    if file_bytes.startswith(b"%PDF"):
        return ".pdf"
    elif file_bytes.startswith(b"PK\x03\x04"):
        return guess_docx_based_extension(file_bytes)
    elif file_bytes.startswith(b"{\\rtf"):
        return ".rtf"
    else:
        raise ValueError("Unsupported or unknown file type")
    
    
def guess_docx_based_extension(file_bytes: bytes) -> str:
    """
    Infers Office file type based on ZIP contents.

    Args:
        file_bytes (bytes): ZIP-based Office document bytes.

    Returns:
        str: Inferred extension.
    """
    try:
        with zipfile.ZipFile(BytesIO(file_bytes)) as z:
            namelist = z.namelist()
            if any(name.startswith("word/") for name in namelist):
                return ".docx"
            elif any(name.startswith("ppt/") for name in namelist):
                return ".pptx"
            elif any(name.startswith("xl/") for name in namelist):
                return ".xlsx"
            else:
                raise ValueError("ZIP-based file has unknown Office structure")
            
    except zipfile.BadZipFile:
        raise ValueError("Invalid ZIP-based Office document")
